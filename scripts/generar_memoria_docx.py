"""Genera memoria/memoria.docx en formato académico (Calibri 12, interlineado 1.5).

Convierte un subset de Markdown:
- # Título 1, ## Título 2, ### Título 3
- Listas con `-` o `1.`
- Tablas pipe `| col | col |`
- Bloques de código ``` ```
- Texto plano y `código inline`
- **negrita**

No es un parser completo de Markdown — solo lo que usa nuestra `memoria.md`.

Uso:
    python scripts/generar_memoria_docx.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor


PROYECTO_DIR = Path(__file__).resolve().parents[1]
MD_PATH = PROYECTO_DIR / "memoria" / "memoria.md"
DOCX_PATH = PROYECTO_DIR / "memoria" / "memoria.docx"


# ---------------------------------------------------------------------
# Estilos del documento — Calibri 12, interlineado 1.5
# ---------------------------------------------------------------------

def configurar_estilos(doc: Document) -> None:
    estilos = doc.styles
    normal = estilos["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for nivel, tam in ((1, 18), (2, 14), (3, 12)):
        nombre = f"Heading {nivel}"
        if nombre in [s.name for s in estilos]:
            est = estilos[nombre]
            est.font.name = "Calibri"
            est.font.size = Pt(tam)
            est.font.bold = True

    if "Code" not in [s.name for s in estilos]:
        code = estilos.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code.font.size = Pt(10)
        code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        code.paragraph_format.left_indent = Pt(12)
        code.paragraph_format.space_after = Pt(0)


# ---------------------------------------------------------------------
# Helpers de inline (negrita, código)
# ---------------------------------------------------------------------

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def _agregar_runs(parrafo, texto: str) -> None:
    pos = 0
    for m in INLINE_RE.finditer(texto):
        if m.start() > pos:
            parrafo.add_run(texto[pos : m.start()])
        token = m.group(0)
        if token.startswith("**"):
            r = parrafo.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("`"):
            r = parrafo.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        pos = m.end()
    if pos < len(texto):
        parrafo.add_run(texto[pos:])


# ---------------------------------------------------------------------
# Conversor línea-a-línea
# ---------------------------------------------------------------------

def convertir(md: str, doc: Document) -> None:
    lineas = md.splitlines()
    i = 0
    n = len(lineas)
    en_codigo = False
    buffer_codigo: list[str] = []

    while i < n:
        linea = lineas[i]

        # Bloque de código ```
        if linea.strip().startswith("```"):
            if not en_codigo:
                en_codigo = True
                buffer_codigo = []
            else:
                # cerrar bloque
                for ln in buffer_codigo:
                    p = doc.add_paragraph(ln, style="Code")
                en_codigo = False
            i += 1
            continue
        if en_codigo:
            buffer_codigo.append(linea)
            i += 1
            continue

        if not linea.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Imagen ![alt](ruta)
        m_img = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", linea)
        if m_img:
            alt, ruta_rel = m_img.group(1), m_img.group(2)
            ruta_abs = (MD_PATH.parent / ruta_rel).resolve()
            if ruta_abs.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(ruta_abs), width=Inches(4.0))
            else:
                p = doc.add_paragraph()
                r = p.add_run(f"[imagen no encontrada: {ruta_rel}]")
                r.italic = True
                r.font.color.rgb = RGBColor(0xAA, 0x33, 0x33)
            i += 1
            continue

        # Encabezados
        if linea.startswith("### "):
            doc.add_heading(linea[4:].strip(), level=3)
            i += 1
            continue
        if linea.startswith("## "):
            doc.add_heading(linea[3:].strip(), level=2)
            i += 1
            continue
        if linea.startswith("# "):
            doc.add_heading(linea[2:].strip(), level=1)
            i += 1
            continue

        # Tabla pipe
        if linea.startswith("|") and i + 1 < n and re.match(r"^\|[\s\-:|]+\|$", lineas[i + 1]):
            filas: list[list[str]] = []
            while i < n and lineas[i].startswith("|"):
                if re.match(r"^\|[\s\-:|]+\|$", lineas[i]):
                    i += 1
                    continue
                celdas = [c.strip() for c in lineas[i].strip().strip("|").split("|")]
                filas.append(celdas)
                i += 1
            if filas:
                tabla = doc.add_table(rows=len(filas), cols=len(filas[0]))
                tabla.style = "Light Grid Accent 1"
                for r_idx, fila in enumerate(filas):
                    celdas = tabla.rows[r_idx].cells
                    for c_idx, valor in enumerate(fila):
                        if c_idx < len(celdas):
                            celdas[c_idx].text = ""
                            p = celdas[c_idx].paragraphs[0]
                            _agregar_runs(p, valor)
                            if r_idx == 0:
                                for run in p.runs:
                                    run.bold = True
            continue

        # Citas / blockquote (>)
        if linea.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            r = p.add_run(linea[2:])
            r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Lista numerada
        m = re.match(r"^\s*\d+\.\s+(.+)$", linea)
        if m:
            p = doc.add_paragraph(style="List Number")
            _agregar_runs(p, m.group(1))
            i += 1
            continue

        # Lista con viñeta
        m = re.match(r"^\s*[-*]\s+(.+)$", linea)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _agregar_runs(p, m.group(1))
            i += 1
            continue

        # Línea horizontal
        if linea.strip() == "---":
            doc.add_paragraph().add_run("─" * 60)
            i += 1
            continue

        # Párrafo normal
        p = doc.add_paragraph()
        _agregar_runs(p, linea)
        i += 1


def main() -> int:
    if not MD_PATH.exists():
        print(f"No existe {MD_PATH}", file=sys.stderr)
        return 1

    md = MD_PATH.read_text(encoding="utf-8")

    doc = Document()
    configurar_estilos(doc)
    convertir(md, doc)
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_PATH))
    print(f"Generado: {DOCX_PATH}  ({DOCX_PATH.stat().st_size / 1024:.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
